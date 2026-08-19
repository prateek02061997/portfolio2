"""
Cover Letter Builder — Stage 2 of the pipeline (deterministic rendering).
Same principle as cv_builder.py: this file has zero LLM calls. It just
renders whatever content dict it's given, the same way every time.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Calibri"
NAME_SIZE = Pt(15)
BODY_SIZE = Pt(11)
MUTED_COLOR = RGBColor(0x44, 0x44, 0x44)


def _set_font(run, size=BODY_SIZE, bold=False, color=None):
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)


def _add_hyperlink(paragraph, url, text, size=Pt(9.5)):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size.pt * 2)))
    rPr.append(sz)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1155CC')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _para(doc, text=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10,
          bold=False, size=BODY_SIZE, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold, color=color)
    return p


def build_cover_letter(data: dict, output_path: str):
    """
    data schema:
    {
        "name": "Prateek Parihar",
        "contact_line": "Auckland, NZ | +64 ... | email",
        "links": [{"text": "...", "url": "..."}, ...],   # optional
        "date": "25 July 2026",
        "recipient_lines": ["Hiring Team", "Company Name", "Location"],
        "salutation": "Dear Hiring Team,",
        "body_paragraphs": ["para 1 text", "para 2 text", ...],
        "closing": "Kind regards,",
        "sign_off_name": "Prateek Parihar"
    }
    """
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.76)
    section.bottom_margin = Cm(1.76)
    section.left_margin = Cm(1.94)
    section.right_margin = Cm(1.94)

    _para(doc, data["name"].upper(), align=WD_ALIGN_PARAGRAPH.CENTER,
          space_after=3, bold=True, size=NAME_SIZE)

    _para(doc, data["contact_line"], align=WD_ALIGN_PARAGRAPH.CENTER,
          space_after=3, size=Pt(9.5), color=MUTED_COLOR)

    if data.get("links"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(20)
        for i, link_obj in enumerate(data["links"]):
            if i > 0:
                sep = p.add_run(" | ")
                _set_font(sep, size=Pt(9.5), color=MUTED_COLOR)
            _add_hyperlink(p, link_obj["url"], link_obj["text"])

    _para(doc, data["date"])
    for line in data["recipient_lines"]:
        _para(doc, line)

    _para(doc, data["salutation"])

    for para_text in data["body_paragraphs"]:
        _para(doc, para_text)

    _para(doc, data["closing"], space_after=30)
    _para(doc, data["sign_off_name"])

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    import json
    import sys

    content_path = sys.argv[1] if len(sys.argv) > 1 else "cover_letter_content_example.json"
    output_path = sys.argv[2] if len(sys.argv) > 2 else "output_cover_letter.docx"

    with open(content_path) as f:
        data = json.load(f)

    build_cover_letter(data, output_path)
    print(f"Saved: {output_path}")
