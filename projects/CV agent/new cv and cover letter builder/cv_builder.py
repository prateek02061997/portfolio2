"""
CV Builder — Stage 2 of the pipeline (deterministic rendering).

This script takes STRUCTURED CONTENT (a Python dict / JSON) and renders it into
a formatted .docx file. It contains ZERO calls to any LLM. All fonts, colors,
sizes, and spacing are fixed constants below, so output is 100% consistent
every single time you run it.

Stage 1 (a separate Claude API call) is responsible ONLY for producing the
JSON content that gets passed into build_cv(). See claude_stage1_prompt.md
and content_schema_example.json for that half of the pipeline.

Usage:
    from cv_builder import build_cv
    import json

    with open("content.json") as f:
        data = json.load(f)

    build_cv(data, "output/My_CV.docx")
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# FIXED STYLE CONSTANTS — change these once, they apply everywhere.
# Never let an LLM choose these values; that's how formatting drifts.
# ---------------------------------------------------------------------------
FONT = "Calibri"
NAME_SIZE = Pt(16)
HEAD_SIZE = Pt(11)
BODY_SIZE = Pt(10.5)
HEADING_COLOR = RGBColor(0x1F, 0x1F, 0x1F)
LINK_COLOR = RGBColor(0x11, 0x55, 0xCC)
MUTED_COLOR = RGBColor(0x44, 0x44, 0x44)
MARGIN_TOP_BOTTOM = Cm(1.24)   # ~700 twips
MARGIN_LEFT_RIGHT = Cm(1.59)   # ~900 twips


def _set_font(run, size=BODY_SIZE, bold=False, color=None, underline=False):
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    run.font.underline = underline
    if color:
        run.font.color.rgb = color
    # Ensure east-asian font fallback doesn't override Calibri
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)


def _add_hyperlink(paragraph, url, text, size=BODY_SIZE):
    """Insert a real, clickable hyperlink run into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size.pt * 2)))
    rPr.append(sz)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1155CC')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(5)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '444444')
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(text.upper())
    _set_font(run, size=HEAD_SIZE, bold=True, color=HEADING_COLOR)
    return p


def _job_header(doc, title, dates):
    """Title on the left, dates right-aligned via a right tab stop."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17.0), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(title)
    _set_font(r1, bold=True)
    r2 = p.add_run("\t" + dates)
    _set_font(r2, bold=True)
    return p


def _project_header(doc, title, link_text=None, link_url=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(title + ("  " if link_text else ""))
    _set_font(r1, bold=True)
    if link_text and link_url:
        _add_hyperlink(p, link_url, link_text)
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    _set_font(run)
    return p


def _plain(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _set_font(run, color=color)
    return p


def build_cv(data: dict, output_path: str):
    """
    data: dict matching content_schema_example.json
    output_path: where to save the .docx
    """
    doc = Document()
    section = doc.sections[0]
    section.top_margin = MARGIN_TOP_BOTTOM
    section.bottom_margin = MARGIN_TOP_BOTTOM
    section.left_margin = MARGIN_LEFT_RIGHT
    section.right_margin = MARGIN_LEFT_RIGHT

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(data["name"].upper())
    _set_font(run, size=NAME_SIZE, bold=True)

    # Contact line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    contact_text = " | ".join(data["contact_line"])
    run = p.add_run(contact_text)
    _set_font(run)

    # Links line (all hyperlinks, clickable)
    if data.get("links"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        for i, link_obj in enumerate(data["links"]):
            if i > 0:
                sep = p.add_run(" | ")
                _set_font(sep)
            _add_hyperlink(p, link_obj["url"], link_obj["text"])

    # Visa / tagline line
    if data.get("tagline"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(data["tagline"])
        _set_font(run, color=MUTED_COLOR)

    # Summary
    _section_heading(doc, "Professional Summary")
    _plain(doc, data["summary"])

    # Skills
    _section_heading(doc, "Skills")
    for line in data["skills"]:
        _plain(doc, line)

    # Projects
    _section_heading(doc, "Analytics Projects")
    for proj in data["projects"]:
        _project_header(doc, proj["title"], proj.get("link_text"), proj.get("link_url"))
        for b in proj["bullets"]:
            _bullet(doc, b)

    # --- PAGE BREAK before Experience: explicit, not relying on natural flow ---
    doc.add_page_break()

    # Experience
    _section_heading(doc, "Professional Experience")
    for job in data["experience"]:
        _job_header(doc, job["title"], job["dates"])
        for b in job["bullets"]:
            _bullet(doc, b)

    # Education
    _section_heading(doc, "Education")
    for edu in data["education"]:
        _job_header(doc, edu["title"], edu["dates"])
        for b in edu.get("bullets", []):
            _bullet(doc, b)

    # Certifications
    _section_heading(doc, "Certifications")
    _plain(doc, data["certifications"])

    # Additional
    _section_heading(doc, "Additional Experience")
    for line in data["additional"]:
        _plain(doc, line)

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    import json
    import sys

    content_path = sys.argv[1] if len(sys.argv) > 1 else "content_schema_example.json"
    output_path = sys.argv[2] if len(sys.argv) > 2 else "output_cv.docx"

    with open(content_path) as f:
        data = json.load(f)

    build_cv(data, output_path)
    print(f"Saved: {output_path}")
